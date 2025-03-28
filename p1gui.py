import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import xlsxwriter
from operator import itemgetter
import os

def process_files(input_file, output_file):
    """
    This function encapsulates your original logic:
    1. Opens a docx file
    2. Underlines repeated lines
    3. Saves an updated docx
    4. Creates an Excel file with frequency counts
    """
    # Read the input .docx file
    with open(input_file, "rb") as f:
        doc_in = Document(f)
        duplist = [para.text for para in doc_in.paragraphs]

    # Prepare a new Document for output
    doc_out = Document()
    doc_out.styles["Normal"].paragraph_format.space_after = 0

    # Dictionary to count frequencies
    freq_count = {}

    # Underline repeated lines
    underline_next = False
    for i, phrase in enumerate(duplist):
        # Update frequency
        freq_count[phrase] = freq_count.get(phrase, 0) + 1

        # Determine if current phrase is redundant (identical to previous line)
        if i > 0 and phrase == duplist[i - 1]:
            underline_next = True

        # Write the line to the output Document
        paragraph = doc_out.add_paragraph()
        run = paragraph.add_run(phrase)
        if underline_next:
            run.underline = True
            underline_next = False

    # Save the underlined .docx file
    doc_out.save(output_file)

    # Sort the dictionary by descending frequency
    sorted_frequencies = sorted(freq_count.items(), key=itemgetter(1), reverse=True)

    # Create Excel sheet with frequency counts
    workbook = xlsxwriter.Workbook("frequencyTable.xlsx")
    worksheet = workbook.add_worksheet()

    for row_idx, (text, count) in enumerate(sorted_frequencies):
        worksheet.write(row_idx, 0, text)
        worksheet.write(row_idx, 1, count)

    workbook.close()


def select_input_file():
    """
    Prompt user to select an input .docx file.
    """
    file_path = filedialog.askopenfilename(
        title="Select Input .docx File",
        filetypes=[("DOCX files", "*.docx"), ("All files", "*.*")]
    )
    if file_path:
        input_entry.delete(0, tk.END)
        input_entry.insert(0, file_path)


def select_output_file():
    """
    Prompt user to select (or name) an output .docx file.
    """
    file_path = filedialog.asksaveasfilename(
        title="Save Output .docx As",
        defaultextension=".docx",
        filetypes=[("DOCX files", "*.docx"), ("All files", "*.*")]
    )
    if file_path:
        output_entry.delete(0, tk.END)
        output_entry.insert(0, file_path)


def run_process():
    """
    Validate entries and run the main docx processing.
    """
    in_file = input_entry.get().strip()
    out_file = output_entry.get().strip()

    if not in_file or not out_file:
        messagebox.showerror("Error", "Please select both input and output files.")
        return

    try:
        process_files(in_file, out_file)
        messagebox.showinfo("Done", f"Output saved to:\n{out_file}\n\nFrequency table: frequencyTable.xlsx")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred:\n{e}")


# ---------------------------
# TKINTER GUI SETUP
# ---------------------------
root = tk.Tk()
root.title("DOCX Redundancy Checker")

# Input file selection
input_frame = tk.Frame(root)
input_frame.pack(pady=5, padx=10, fill="x")

input_label = tk.Label(input_frame, text="Input .docx:")
input_label.pack(side="left")

input_entry = tk.Entry(input_frame, width=50)
input_entry.pack(side="left", padx=5)

input_button = tk.Button(input_frame, text="Browse", command=select_input_file)
input_button.pack(side="left")

# Output file selection
output_frame = tk.Frame(root)
output_frame.pack(pady=5, padx=10, fill="x")

output_label = tk.Label(output_frame, text="Output .docx:")
output_label.pack(side="left")

output_entry = tk.Entry(output_frame, width=50)
output_entry.pack(side="left", padx=5)

output_button = tk.Button(output_frame, text="Browse", command=select_output_file)
output_button.pack(side="left")

# Run button
run_button = tk.Button(root, text="Process Files", command=run_process)
run_button.pack(pady=15)

root.mainloop()
